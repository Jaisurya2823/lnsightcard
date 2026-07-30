"""Tests for the Excel, Word, and plain-text exporters — each checked
for producing a valid, openable file with the expected structure."""

from __future__ import annotations

import pandas as pd
import pytest
from docx import Document
from openpyxl import load_workbook

from insightcard import data_quality, docx_report, excel_report, metrics as metrics_mod, txt_report
from insightcard.config import BrandConfig


def _extract_pdf_text(pdf_path) -> str:
    import shutil
    import subprocess

    if shutil.which("pdftotext") is None:
        pytest.skip("pdftotext (poppler-utils) not installed -- skipping PDF text-content check")
    result = subprocess.run(["pdftotext", "-layout", str(pdf_path), "-"], capture_output=True, text=True)
    return result.stdout


@pytest.fixture
def sample_metrics():
    df = pd.DataFrame(
        {
            "date": pd.to_datetime(
                ["2024-01-01", "2024-02-01", "2024-03-01", "2024-04-01", "2024-05-01"]
            ),
            "product": ["Mug", "Bottle", "Mug", "Plate", "Bottle"],
            "category": ["Home", "Home", "Home", "Kitchen", "Home"],
            "amount": [500.0, 300.0, 600.0, 200.0, 350.0],
        }
    )
    m = metrics_mod.compute_metrics(df)
    qr = data_quality.check_data_quality(df)
    return m, qr


def test_excel_workbook_has_expected_sheets(sample_metrics, tmp_path):
    m, qr = sample_metrics
    brand = BrandConfig(business_name="Test Shop")
    out = excel_report.build_workbook(m, brand, tmp_path / "report.xlsx", data_quality=qr)
    assert out.exists()
    wb = load_workbook(out)
    assert "Summary" in wb.sheetnames
    assert "Top Products" in wb.sheetnames
    assert "Monthly Revenue" in wb.sheetnames


def test_excel_summary_sheet_contains_kpis(sample_metrics, tmp_path):
    m, qr = sample_metrics
    brand = BrandConfig(business_name="Test Shop", currency_symbol="Rs ")
    out = excel_report.build_workbook(m, brand, tmp_path / "report.xlsx", data_quality=qr)
    wb = load_workbook(out)
    ws = wb["Summary"]
    all_values = [str(cell.value) for row in ws.iter_rows() for cell in row if cell.value is not None]
    joined = " ".join(all_values)
    assert "Test Shop" in joined
    assert f"{m.total_revenue:,.2f}" in joined


def test_docx_has_title_and_table(sample_metrics, tmp_path):
    m, qr = sample_metrics
    brand = BrandConfig(business_name="Test Shop")
    out = docx_report.build_docx(m, [], brand, tmp_path / "report.docx", data_quality=qr)
    assert out.exists()
    doc = Document(out)
    assert doc.paragraphs[0].text == brand.report_title
    assert len(doc.tables) >= 1  # at least the Top Products table


def test_docx_includes_insights(sample_metrics, tmp_path):
    m, qr = sample_metrics
    brand = BrandConfig(business_name="Test Shop")
    out = docx_report.build_docx(m, [], brand, tmp_path / "report.docx", data_quality=qr)
    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert m.insights[0] in full_text


def test_txt_report_is_plain_and_contains_kpis(sample_metrics, tmp_path):
    m, qr = sample_metrics
    brand = BrandConfig(business_name="Test Shop", currency_symbol="Rs ")
    out = txt_report.build_txt(m, brand, tmp_path / "report.txt", data_quality=qr)
    assert out.exists()
    content = out.read_text(encoding="utf-8")
    assert "Test Shop" in content
    assert f"{m.total_revenue:,.2f}" in content
    assert "KEY INSIGHTS" in content


def test_txt_report_handles_no_data_quality_arg(sample_metrics, tmp_path):
    m, _ = sample_metrics
    brand = BrandConfig(business_name="Test Shop")
    out = txt_report.build_txt(m, brand, tmp_path / "report.txt", data_quality=None)
    assert out.exists()
    assert "DATA HEALTH CHECK" not in out.read_text(encoding="utf-8")


def test_pdf_renders_accented_unicode_correctly(tmp_path):
    """Regression test: the PDF used to force every string through
    latin-1 encoding (errors='replace'), which silently mangled any
    accented character, currency symbol, or non-English text into '?'.
    Fixed by embedding DejaVu Sans instead of using the core Helvetica
    font. This checks the accented business name and a non-ASCII
    currency symbol actually appear in the rendered PDF text layer."""
    from insightcard import pdf_report

    df = pd.DataFrame({
        "amount": [500.0, 700.0],
        "date": pd.to_datetime(["2024-01-01", "2024-02-01"]),
    })
    m = metrics_mod.compute_metrics(df)
    brand = BrandConfig(business_name="Nöel's Café & Boutique", currency_symbol="€")
    out = pdf_report.build_report(m, [], brand, tmp_path / "report.pdf")
    assert out.exists()
    text = _extract_pdf_text(out)
    assert "Nöel's Café & Boutique" in text
    assert "€" in text

