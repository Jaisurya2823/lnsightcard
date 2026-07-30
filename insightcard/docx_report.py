"""
Word (.docx) exporter — mirrors the PDF report's content, for
businesses that want to further edit the report themselves (add a
note, adjust wording) before sharing it, rather than being stuck with
a flattened PDF. Reuses the same chart PNGs the PDF uses.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from insightcard.config import BrandConfig
from insightcard.metrics import ReportMetrics

_ACCENT = RGBColor(0x1A, 0x23, 0x7E)
_GREEN = RGBColor(0x2E, 0x7D, 0x32)
_RED = RGBColor(0xC6, 0x28, 0x28)
_GREY = RGBColor(0x45, 0x5A, 0x64)


def _heading(doc: Document, text: str, level: int = 1, color: RGBColor = _ACCENT) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color


def _kpi_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(value)


def build_docx(
    metrics: ReportMetrics,
    chart_paths: list[Path],
    brand: BrandConfig,
    output_path: str | Path,
    data_quality=None,
) -> Path:
    doc = Document()

    title = doc.add_heading(brand.report_title, level=0)
    for run in title.runs:
        run.font.color.rgb = _ACCENT

    sub = doc.add_paragraph()
    sub.add_run(brand.business_name).bold = True
    if metrics.date_range:
        doc.add_paragraph(f"Period covered: {metrics.date_range[0]} to {metrics.date_range[1]}")

    # ------------------------------------------------------------- KPIs
    _heading(doc, "Key Metrics", level=2)
    _kpi_line(doc, "Total Revenue", f"{brand.currency_symbol}{metrics.total_revenue:,.2f}")
    _kpi_line(doc, "Orders", f"{metrics.order_count:,}")
    _kpi_line(doc, "Average Order Value", f"{brand.currency_symbol}{metrics.average_order_value:,.2f}")
    if metrics.growth_pct is not None:
        _kpi_line(doc, "Growth (last period)", f"{metrics.growth_pct:+.1f}%")

    # ------------------------------------------------- Data Health Check
    if data_quality is not None:
        _heading(doc, "Data Health Check", level=2)
        if data_quality.is_clean:
            p = doc.add_paragraph(f"No data quality issues detected across {data_quality.total_rows:,} rows.")
            p.runs[0].font.color.rgb = _GREEN
        else:
            for warning in data_quality.warnings:
                doc.add_paragraph(warning, style="List Bullet")

    # -------------------------------------------------------- Insights
    _heading(doc, "Key Insights", level=2)
    for insight in metrics.insights:
        doc.add_paragraph(insight, style="List Bullet")

    # ---------------------------------------------------- Top Products
    if metrics.top_products is not None and len(metrics.top_products):
        _heading(doc, "Top Products", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = "Product", "Revenue"
        for name, revenue in metrics.top_products.items():
            row = table.add_row().cells
            row[0].text = str(name)
            row[1].text = f"{brand.currency_symbol}{revenue:,.2f}"

    # ----------------------------------------------------- Slow Moving
    if metrics.slow_moving_products is not None and len(metrics.slow_moving_products):
        _heading(doc, "Slow-Moving Stock (no sale in 30+ days)", level=2, color=_RED)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 2"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = "Product", "Days Since Last Sale"
        for name, days in metrics.slow_moving_products.head(10).items():
            row = table.add_row().cells
            row[0].text = str(name)
            row[1].text = str(int(days))

    # ------------------------------------------------------------ Charts
    for chart_path in chart_paths:
        title_text = chart_path.stem.replace("chart_", "").replace("_", " ").title()
        _heading(doc, title_text, level=2)
        doc.add_picture(str(chart_path), width=Inches(6))

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(brand.footer_text)
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = _GREY

    output_path = Path(output_path)
    doc.save(str(output_path))
    return output_path
